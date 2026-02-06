import io
import logging
from abc import ABC, abstractmethod
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from docx import Document as DocxDocument
from minio import Minio
from minio.error import S3Error
from pypdf import PdfReader

logger = logging.getLogger(__name__)


@dataclass
class LoadedDocument:
    """Đại diện cho một document đã được load"""
    content: str
    source: str
    metadata: dict
    
    def __len__(self) -> int:
        """Độ dài content"""
        return len(self.content)
    
    def is_empty(self) -> bool:
        """Kiểm tra document có rỗng không"""
        return len(self.content.strip()) == 0


class BaseDocumentLoader(ABC):
    """Base class cho document loaders"""
    
    @abstractmethod
    def load(self) -> List[LoadedDocument]:
        """
        Load documents
        
        Returns:
            List[LoadedDocument]: Danh sách documents
        """
        pass
    
    def _read_docx(self, data: bytes) -> str:
        """
        Đọc DOCX file từ bytes
        
        Args:
            data: Binary data của DOCX file
            
        Returns:
            str: Text content
        """
        try:
            doc = DocxDocument(io.BytesIO(data))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    
    def _read_pdf(self, data: bytes) -> str:
        """
        Đọc PDF file từ bytes
        
        Args:
            data: Binary data của PDF file
            
        Returns:
            str: Text content
        """
        try:
            reader = PdfReader(io.BytesIO(data))
            pages = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n".join(pages)
        except Exception as e:
            logger.error(f"Error reading PDF: {e}")
            raise


class MinIODocumentLoader(BaseDocumentLoader):
    """
    Loader cho documents từ MinIO storage.
    
    Cải tiến:
    - Error handling tốt hơn
    - Logging chi tiết
    - Retry logic
    - Progress tracking
    - Filter theo file types
    """
    
    SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
    
    def __init__(
        self,
        endpoint: str,
        access_key: str,
        secret_key: str,
        bucket: str,
        prefix: str = "",
        secure: bool = False,
        max_retries: int = 3,
    ):
        """
        Initialize MinIO loader
        
        Args:
            endpoint: MinIO endpoint
            access_key: Access key
            secret_key: Secret key
            bucket: Bucket name
            prefix: Object prefix/folder
            secure: Use HTTPS
            max_retries: Max retry attempts
        """
        self.endpoint = endpoint
        self.bucket = bucket
        self.prefix = prefix
        self.max_retries = max_retries
        
        try:
            self.client = Minio(
                endpoint,
                access_key=access_key,
                secret_key=secret_key,
                secure=secure,
            )
            logger.info(f"Connected to MinIO at {endpoint}")
        except Exception as e:
            logger.error(f"Failed to connect to MinIO: {e}")
            raise
    
    def _is_supported_file(self, filename: str) -> bool:
        """Kiểm tra file có được hỗ trợ không"""
        return Path(filename).suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def _get_file_type(self, filename: str) -> Optional[str]:
        """Lấy loại file"""
        suffix = Path(filename).suffix.lower()
        if suffix == ".docx":
            return "docx"
        elif suffix == ".pdf":
            return "pdf"
        return None
    
    def _load_object(self, object_name: str) -> Optional[LoadedDocument]:
        """
        Load một object từ MinIO
        
        Args:
            object_name: Tên object
            
        Returns:
            LoadedDocument hoặc None nếu fail
        """
        file_type = self._get_file_type(object_name)
        if not file_type:
            logger.warning(f"Unsupported file type: {object_name}")
            return None
        
        # Retry logic
        for attempt in range(self.max_retries):
            try:
                logger.info(f"Loading {object_name} (attempt {attempt + 1}/{self.max_retries})")
                
                # Get object data
                response = self.client.get_object(self.bucket, object_name)
                data = response.read()
                response.close()
                response.release_conn()
                
                # Extract text based on file type
                if file_type == "docx":
                    text = self._read_docx(data)
                elif file_type == "pdf":
                    text = self._read_pdf(data)
                else:
                    logger.warning(f"Unknown file type: {file_type}")
                    return None
                
                # Create document
                doc = LoadedDocument(
                    content=text,
                    source=object_name,
                    metadata={
                        "file_type": file_type,
                        "bucket": self.bucket,
                        "size": len(data),
                    }
                )
                
                if doc.is_empty():
                    logger.warning(f"Empty document: {object_name}")
                    return None
                
                logger.info(f"Successfully loaded {object_name} ({len(doc)} chars)")
                return doc
                
            except S3Error as e:
                logger.error(f"S3 error loading {object_name}: {e}")
                if attempt == self.max_retries - 1:
                    raise
            except Exception as e:
                logger.error(f"Error loading {object_name}: {e}")
                if attempt == self.max_retries - 1:
                    raise
        
        return None
    
    def load(self) -> List[LoadedDocument]:
        """
        Load tất cả documents từ MinIO bucket
        
        Returns:
            List[LoadedDocument]: Danh sách documents
        """
        documents = []
        
        try:
            # Check if bucket exists
            if not self.client.bucket_exists(self.bucket):
                raise ValueError(f"Bucket '{self.bucket}' does not exist")
            
            logger.info(f"Listing objects in bucket '{self.bucket}' with prefix '{self.prefix}'")
            
            # List objects
            objects = self.client.list_objects(
                self.bucket,
                prefix=self.prefix,
                recursive=True,
            )
            
            # Load each object
            loaded_count = 0
            skipped_count = 0
            
            for obj in objects:
                object_name = obj.object_name
                
                # Skip directories
                if object_name.endswith('/'):
                    continue
                
                # Check if supported
                if not self._is_supported_file(object_name):
                    logger.debug(f"Skipping unsupported file: {object_name}")
                    skipped_count += 1
                    continue
                
                # Load document
                doc = self._load_object(object_name)
                if doc:
                    documents.append(doc)
                    loaded_count += 1
                else:
                    skipped_count += 1
            
            logger.info(
                f"Loaded {loaded_count} documents, skipped {skipped_count} files"
            )
            
        except S3Error as e:
            logger.error(f"MinIO error: {e}")
            raise
        except Exception as e:
            logger.error(f"Unexpected error: {e}")
            raise
        
        return documents
    
    def load_specific_files(self, file_paths: List[str]) -> List[LoadedDocument]:
        """
        Load specific files from MinIO
        
        Args:
            file_paths: List of file paths relative to bucket
            
        Returns:
            List[LoadedDocument]: Loaded documents
        """
        documents = []
        
        for file_path in file_paths:
            full_path = f"{self.prefix}/{file_path}".lstrip('/')
            doc = self._load_object(full_path)
            if doc:
                documents.append(doc)
        
        logger.info(f"Loaded {len(documents)}/{len(file_paths)} specific files")
        return documents


# Factory function
def create_document_loader(
    loader_type: str = "minio",
    **kwargs
) -> BaseDocumentLoader:
    """
    Factory function để tạo document loader
    
    Args:
        loader_type: Loại loader (hiện tại chỉ support 'minio')
        **kwargs: Arguments cho loader
        
    Returns:
        BaseDocumentLoader instance
    """
    if loader_type == "minio":
        return MinIODocumentLoader(**kwargs)
    else:
        raise ValueError(f"Unsupported loader type: {loader_type}")