from concurrent.futures import ThreadPoolExecutor, as_completed
from logger import Logger
import pickle
import os
from docx import Document


class DocumentProcessor:
    """
    Processes documents with:
    - DOCX reading (text + tables)
    - Chunking
    - Entity / relation extraction via LLM
    - Caching

    NOTE:
    - Does NOT write to Neo4j
    - Returns extracted data only
    """

    logger = Logger("DocumentProcessor").get_logger()

    def __init__(self, client, model="gpt-4o-mini", max_workers=10):
        self.client = client
        self.model = model
        self.max_workers = max_workers
        self.logger.info(
            f"Initialized DocumentProcessor with model={model}, workers={max_workers}"
        )

    # =========================================================
    # DOCX LOADING
    # =========================================================
    def extract_docx_content(self, docx_path):
        """
        Extract text + tables from a DOCX file.

        Returns:
            {
                "text": str,
                "tables": list[dict]
            }
        """
        doc = Document(docx_path)

        # ---- TEXT ----
        paragraphs = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        ]

        # ---- TABLES ----
        tables = extract_tables_from_docx(doc)

        self.logger.info(
            f"Extracted {len(paragraphs)} paragraphs and {len(tables)} tables from {docx_path}"
        )

        return {
            "text": "\n".join(paragraphs),
            "tables": tables
        }

    # =========================================================
    # CHUNKING
    # =========================================================
    def split_documents(self, documents, chunk_size=2000, overlap_size=300):
        chunks = []

        for doc_idx, document in enumerate(documents):
            for i in range(0, len(document), chunk_size - overlap_size):
                chunks.append(document[i:i + chunk_size])

        self.logger.info(
            f"Split {len(documents)} documents into {len(chunks)} chunks"
        )
        return chunks

    # =========================================================
    # LLM EXTRACTION
    # =========================================================
    def extract_elements(self, chunks, use_parallel=True):
        self.logger.info(f"Extracting elements from {len(chunks)} chunks...")

        system_prompt = """You are an information extraction system.

Extract ENTITIES and RELATIONSHIPS from the text.

STRICT FORMAT (no explanation, no markdown):

ENTITY: <entity name>
RELATION: <entity_1> -> <relation> -> <entity_2>

Rules:
- Use '->' exactly
- Entity names: max 5 words
- Use Vietnamese if the text is Vietnamese
- Do NOT invent relations not present in text
"""

        if use_parallel:
            results = self._batch_api_call(
                chunks,
                system_prompt,
                user_template="{item}",
                max_tokens=400
            )
        else:
            results = []
            for i, chunk in enumerate(chunks):
                try:
                    response = self.client.chat.completions.create(
                        model=self.model,
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": chunk[:1500]}
                        ],
                        max_tokens=400
                    )
                    results.append(response.choices[0].message.content)
                except Exception as e:
                    self.logger.warning(f"Chunk {i} failed: {e}")
                    results.append("")

        valid = [r for r in results if r]
        self.logger.info(f"Extracted {len(valid)} valid element sets")
        return valid

    def _batch_api_call(self, items, system_prompt, user_template, max_tokens):
        results = [None] * len(items)

        def worker(index, item):
            try:
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_template.format(item=item[:1500])}
                    ],
                    max_tokens=max_tokens
                )
                return index, response.choices[0].message.content
            except Exception as e:
                self.logger.warning(f"Item {index} failed: {e}")
                return index, ""

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            futures = [
                executor.submit(worker, i, item)
                for i, item in enumerate(items)
            ]
            for future in as_completed(futures):
                idx, res = future.result()
                results[idx] = res

        return [r for r in results if r]

    # =========================================================
    # CACHE
    # =========================================================
    def load_or_process(self, file_path, process_function, *args, **kwargs):
        directory = os.path.dirname(file_path)
        if directory and not os.path.exists(directory):
            os.makedirs(directory)

        if os.path.exists(file_path):
            self.logger.info(f"Loading cached data from {file_path}")
            with open(file_path, "rb") as f:
                return pickle.load(f)

        self.logger.info(f"Cache miss - running process for {file_path}")
        data = process_function(*args, **kwargs)

        if data is not None:
            with open(file_path, "wb") as f:
                pickle.dump(data, f)

        return data


# =========================================================
# DOCX DIRECTORY LOADER (USED BY app.py)
# =========================================================
def read_docx_from_directory(directory_path):
    """
    Read all DOCX files from a directory and return combined text.

    Returns:
        list[str]
    """
    documents = []

    for filename in os.listdir(directory_path):
        if not filename.lower().endswith(".docx"):
            continue
        if filename.startswith("~"):
            continue

        path = os.path.join(directory_path, filename)
        doc = Document(path)

        texts = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text.strip()
        ]

        if texts:
            documents.append("\n".join(texts))

    return documents


# =========================================================
# TABLE EXTRACTION (INTERNAL)
# =========================================================
def get_cell_text(cell):
    return "\n".join(
        p.text.strip()
        for p in cell.paragraphs
        if p.text.strip()
    )


def get_gridspan(cell):
    tc = cell._tc
    gs = tc.tcPr.gridSpan
    return int(gs.val) if gs is not None else 1


def get_vmerge(cell):
    tc = cell._tc
    vm = tc.tcPr.vMerge
    return None if vm is None else vm.val


def is_meaningful_table(headers):
    if len(headers) < 2:
        return False

    blacklist = [
        "bộ giáo dục",
        "cộng hòa",
        "hiệu trưởng",
        "trưởng khoa"
    ]

    joined = " ".join(headers).lower()
    return not any(b in joined for b in blacklist)


def extract_table(table, table_index):
    matrix = []
    rowspan_tracker = {}
    current_row_label = None

    for r, row in enumerate(table.rows):
        row_data = []
        c = 0

        for cell in row.cells:
            while (r, c) in rowspan_tracker:
                row_data.append(rowspan_tracker[(r, c)])
                c += 1

            text = get_cell_text(cell)
            colspan = get_gridspan(cell)
            vmerge = get_vmerge(cell)

            cell_data = {"text": text}

            if vmerge == "restart":
                rowspan_tracker[(r, c)] = cell_data
            elif vmerge is None and (r - 1, c) in rowspan_tracker:
                c += colspan
                continue

            for _ in range(colspan):
                row_data.append(cell_data)
                c += 1

        if row_data:
            matrix.append(row_data)

    if not matrix:
        return None

    headers = [c["text"] for c in matrix[0]]
    rows = []

    for row in matrix[1:]:
        label = row[0]["text"].strip()
        if label:
            current_row_label = label

        row_obj = {
            "row_label": current_row_label,
            "cells": {}
        }

        for i in range(1, min(len(headers), len(row))):
            key = headers[i].strip()
            val = row[i]["text"].strip()
            if key:
                row_obj["cells"][key] = val

        if row_obj["cells"]:
            rows.append(row_obj)

    return {
        "table_index": table_index,
        "headers": headers,
        "rows": rows
    }


def extract_tables_from_docx(doc):
    tables = []

    for idx, table in enumerate(doc.tables):
        table_json = extract_table(table, idx)
        if table_json and is_meaningful_table(table_json["headers"]):
            tables.append(table_json)

    return tables
